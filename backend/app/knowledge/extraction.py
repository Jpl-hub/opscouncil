from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_KNOWLEDGE_UPLOAD_BYTES = 8 * 1024 * 1024
MAX_EXTRACTED_CHARS = 100_000
TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".log",
    ".conf",
    ".cfg",
    ".ini",
    ".service",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
}
OCR_REQUIRED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


class KnowledgeFileRejectedError(ValueError):
    pass


@dataclass(frozen=True)
class ExtractedKnowledgeFile:
    title: str
    source_uri: str
    content: str
    file_type: str
    char_count: int


def extract_knowledge_file(filename: str, data: bytes) -> ExtractedKnowledgeFile:
    safe_name = _safe_filename(filename)
    if not data:
        raise KnowledgeFileRejectedError("上传文件为空。")
    if len(data) > MAX_KNOWLEDGE_UPLOAD_BYTES:
        raise KnowledgeFileRejectedError("上传文件超过 8 MB，请拆分后再入库。")
    suffix = Path(safe_name).suffix.lower()
    if suffix == ".pdf":
        content = _extract_pdf(data)
    elif suffix == ".docx":
        content = _extract_docx(data)
    elif suffix in TEXT_EXTENSIONS:
        content = _extract_text(data)
    elif suffix in OCR_REQUIRED_EXTENSIONS:
        raise KnowledgeFileRejectedError("当前文件疑似图片或扫描件，需要 OCR 通道后才能入库。")
    else:
        raise KnowledgeFileRejectedError("暂不支持该文件格式，请上传 PDF、DOCX、TXT、Markdown 或日志文本。")
    normalized = _normalize_extracted_text(content)
    if len(normalized) < 20:
        raise KnowledgeFileRejectedError("未抽取到足够正文，可能是扫描件 PDF 或空文档。")
    if len(normalized) > MAX_EXTRACTED_CHARS:
        normalized = normalized[:MAX_EXTRACTED_CHARS]
    return ExtractedKnowledgeFile(
        title=Path(safe_name).stem[:256] or "未命名资料",
        source_uri=f"upload://{safe_name}",
        content=normalized,
        file_type=suffix.lstrip(".") or "text",
        char_count=len(normalized),
    )


def _extract_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages[:80]]
    except Exception as exc:  # pypdf raises multiple parser-specific exceptions
        raise KnowledgeFileRejectedError(f"PDF 正文抽取失败：{exc}") from exc
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise KnowledgeFileRejectedError(f"DOCX 正文抽取失败：{exc}") from exc
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise KnowledgeFileRejectedError("文本编码无法识别，请转换为 UTF-8 后再上传。")


def _normalize_extracted_text(content: str) -> str:
    text = content.replace("\x00", " ").strip()
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _safe_filename(filename: str) -> str:
    name = Path(filename or "knowledge.txt").name.strip()
    return re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", name)[:180] or "knowledge.txt"
