from __future__ import annotations

from io import BytesIO
import unittest

from docx import Document

from backend.app.knowledge.extraction import KnowledgeFileRejectedError, extract_knowledge_file


class KnowledgeFileExtractionTest(unittest.TestCase):
    def test_extracts_utf8_text_file(self) -> None:
        extracted = extract_knowledge_file(
            "数据库日志处置.md",
            "数据库事务日志不得自动删除。\n普通应用日志应先备份、压缩并截断。".encode("utf-8"),
        )

        self.assertEqual(extracted.title, "数据库日志处置")
        self.assertEqual(extracted.file_type, "md")
        self.assertIn("数据库事务日志不得自动删除", extracted.content)
        self.assertEqual(extracted.source_uri, "upload://数据库日志处置.md")

    def test_extracts_docx_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_paragraph("SSH 暴露面排查前必须确认堡垒机和业务入口。")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "动作"
        table.cell(0, 1).text = "审批后收敛监听地址"
        buffer = BytesIO()
        document.save(buffer)

        extracted = extract_knowledge_file("ssh-runbook.docx", buffer.getvalue())

        self.assertEqual(extracted.file_type, "docx")
        self.assertIn("SSH 暴露面排查", extracted.content)
        self.assertIn("动作 | 审批后收敛监听地址", extracted.content)

    def test_rejects_image_until_ocr_channel_is_available(self) -> None:
        with self.assertRaises(KnowledgeFileRejectedError) as raised:
            extract_knowledge_file("scan.png", b"not-a-real-image-but-extension-is-image")

        self.assertIn("OCR", str(raised.exception))


if __name__ == "__main__":
    unittest.main()
